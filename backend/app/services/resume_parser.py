"""Extract structured profile fields from an uploaded resume."""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

KNOWN_SKILLS = [
    "Python",
    "JavaScript",
    "TypeScript",
    "Java",
    "SQL",
    "React",
    "Node.js",
    "AWS",
    "Excel",
    "MS Excel",
    "Power BI",
    "Tally",
    "GST",
    "SAP",
    "SAP FICO",
    "QuickBooks",
    "Taxation",
    "Accounts Payable",
    "Accounts Receivable",
    "Bank Reconciliation",
    "Financial Reporting",
    "CRM",
    "Salesforce",
    "Zendesk",
    "Customer Service",
    "Communication",
    "Call Handling",
    "Email Support",
    "Inventory Management",
    "Procurement",
    "Vendor Management",
    "Merchandising",
    "Supply Chain",
    "MS Office",
    "Negotiation",
    "Lead Generation",
    "Recruitment",
    "HRMS",
    "Payroll",
    "MIS Reporting",
    "LangGraph",
    "LangChain",
    "Python",
    "FastAPI",
    "Django",
    "PostgreSQL",
    "MongoDB",
    "Docker",
]

SECTION_ALIASES = {
    "summary": ("summary", "profile", "about", "objective", "career objective", "professional summary"),
    "skills": (
        "technical skills",
        "skills",
        "core competencies",
        "tech stack",
        "key skills",
        "functional skills",
        "computer proficiency",
        "tools known",
        "software skills",
    ),
    "experience": (
        "work experience",
        "professional experience",
        "experience",
        "employment",
        "work history",
    ),
    "education": ("education", "academic", "academics", "qualifications"),
}

DEGREE_RE = re.compile(
    r"\b(B\.?Tech|B\.?E\.?|B\.?Sc|B\.?S\.?|B\.?Com|B\.?A\.?|BBA|BCA|"
    r"M\.?Tech|M\.?S\.?|M\.?Sc|M\.?Com|MBA|MCA|CA|CFA|CS|Ph\.?D\.?|"
    r"Bachelor(?:'s)?|Master(?:'s)?|Diploma)\b",
    re.I,
)
PHONE_RE = re.compile(r"(?:\+?\d[\d\s().-]{8,}\d)")
EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)
DATE_RANGE_RE = re.compile(
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4})"
    r"\s*[-–—to]+\s*"
    r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}|Present|Current|Now)",
    re.I,
)
MONTHS = {
    "jan": "01",
    "feb": "02",
    "mar": "03",
    "apr": "04",
    "may": "05",
    "jun": "06",
    "jul": "07",
    "aug": "08",
    "sep": "09",
    "oct": "10",
    "nov": "11",
    "dec": "12",
}


def extract_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(content)
    if suffix == ".docx":
        return _extract_docx(content)
    if suffix == ".doc":
        return _extract_doc_bytes(content)
    raise ValueError("Unsupported resume type")


def parse_resume(filename: str, content: bytes) -> dict[str, Any]:
    text = extract_text(filename, content)
    if not text.strip():
        return {}

    sections = _split_sections(text)
    phone = _first_match(PHONE_RE, text)
    if phone:
        phone = re.sub(r"[^\d+]", "", phone)
        if not phone.startswith("+") and len(phone) == 10:
            phone = "+91" + phone

    return {
        "full_name": _guess_name(text),
        "phone": phone,
        "location": _guess_location(text),
        "bio": _clean_block(sections.get("summary") or _guess_bio(text)),
        "skills": _extract_skills(sections.get("skills") or text),
        "experience": _extract_experience(sections.get("experience") or ""),
        "education": _extract_education(sections.get("education") or text),
    }


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(content: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(content))
    lines = [para.text for para in document.paragraphs if para.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_doc_bytes(content: bytes) -> str:
    try:
        decoded = content.decode("utf-16le", errors="ignore")
    except Exception:
        decoded = content.decode("latin-1", errors="ignore")
    cleaned = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", decoded)
    return re.sub(r"\s{2,}", "\n", cleaned)


def _normalize(line: str) -> str:
    return re.sub(r"\s+", " ", line).strip()


def _is_heading(line: str) -> str | None:
    compact = re.sub(r"[^a-z\s]", "", line.lower()).strip()
    if not compact or len(compact) > 40:
        return None
    for key, aliases in SECTION_ALIASES.items():
        if compact in aliases:
            return key
    return None


def _split_sections(text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current = "header"
    sections[current] = []
    for raw in text.splitlines():
        line = _normalize(raw)
        if not line:
            continue
        heading = _is_heading(line)
        if heading:
            current = heading
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return {key: "\n".join(lines) for key, lines in sections.items() if lines}


def _first_match(pattern: re.Pattern[str], text: str) -> str | None:
    match = pattern.search(text)
    return match.group(0).strip() if match else None


def _guess_name(text: str) -> str:
    for line in text.splitlines():
        candidate = _normalize(line)
        if not candidate or "@" in candidate or PHONE_RE.search(candidate):
            continue
        if _is_heading(candidate):
            continue
        if 2 <= len(candidate.split()) <= 4 and candidate.replace(" ", "").isalpha():
            if candidate.lower() not in {"technical skills", "work experience"}:
                return candidate.title() if candidate.isupper() else candidate
    return ""


def _guess_location(text: str) -> str:
    cities = [
        "Bangalore",
        "Bengaluru",
        "Mumbai",
        "Delhi",
        "New Delhi",
        "Hyderabad",
        "Pune",
        "Chennai",
        "Kolkata",
        "Gurgaon",
        "Gurugram",
        "Noida",
        "Ahmedabad",
        "Jaipur",
        "Remote",
    ]
    header = "\n".join(text.splitlines()[:12])
    for city in cities:
        if re.search(rf"\b{re.escape(city)}\b", header, re.I):
            return city
    match = re.search(r"\b([A-Z][a-z]+),\s*India\b", header)
    return match.group(0) if match else ""


def _guess_bio(text: str) -> str:
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    for para in paragraphs[:4]:
        if len(para) > 80 and not _is_heading(para.splitlines()[0]):
            return _clean_block(para)
    return ""


def _clean_block(text: str) -> str:
    lines = [_normalize(line) for line in text.splitlines() if _normalize(line)]
    return " ".join(lines)[:800]


def _extract_skills(text: str) -> list[str]:
    found: list[str] = []
    lower = text.lower()
    for skill in sorted(KNOWN_SKILLS, key=len, reverse=True):
        if re.search(rf"(?<![A-Za-z]){re.escape(skill)}(?![A-Za-z])", text, re.I) or skill.lower() in lower:
            if skill not in found:
                found.append(skill)

    for line in text.splitlines():
        if ":" in line:
            _, rest = line.split(":", 1)
            for part in re.split(r"[,|/•;]", rest):
                skill = _normalize(part)
                if 1 < len(skill) <= 40 and skill not in found and not skill.isdigit():
                    found.append(skill)
    return found[:40]


def _to_iso_date(value: str) -> str:
    value = value.strip()
    if re.match(r"(present|current|now)", value, re.I):
        return ""
    month_year = re.match(r"([A-Za-z]+)\.?\s+(\d{4})", value)
    if month_year:
        month = MONTHS.get(month_year.group(1)[:3].lower(), "01")
        return f"{month_year.group(2)}-{month}-01"
    year = re.match(r"(\d{4})$", value)
    if year:
        return f"{year.group(1)}-01-01"
    return ""


def _parse_role_line(line: str) -> tuple[str, str] | None:
    if DATE_RANGE_RE.fullmatch(line) or DEGREE_RE.search(line):
        return None
    if " at " in line.lower():
        title, company = re.split(r"\s+at\s+", line, maxsplit=1, flags=re.I)
        return title.strip(" -,|"), company.strip(" -,|")
    if " | " in line:
        left, right = line.split(" | ", 1)
        return left.strip(" -,|"), right.strip(" -,|")
    if re.search(r"\s[-–]\s", line) and not DATE_RANGE_RE.search(line):
        parts = re.split(r"\s[-–]\s", line, maxsplit=1)
        if len(parts) == 2:
            return parts[0].strip(), parts[1].strip()
    if re.search(
        r"\b(engineer|developer|scientist|manager|architect|intern|founder|lead|"
        r"consultant|analyst|executive|associate|officer|representative|"
        r"coordinator|specialist|accountant|merchandiser|recruiter)\b",
        line,
        re.I,
    ) and len(line.split()) <= 8:
        return line, ""
    return None


def _extract_experience(text: str) -> list[dict[str, str]]:
    if not text.strip():
        return []

    entries: list[dict[str, str]] = []
    current: dict[str, str] | None = None
    desc_parts: list[str] = []

    def flush() -> None:
        nonlocal current, desc_parts
        if current and (current.get("title") or current.get("company")):
            current["description"] = " ".join(desc_parts)[:600]
            entries.append(current)
        current = None
        desc_parts = []

    for raw in text.splitlines():
        line = _normalize(raw)
        if not line or _is_heading(line):
            continue
        date_match = DATE_RANGE_RE.search(line)
        role = _parse_role_line(line) if not date_match else None
        if role:
            flush()
            current = {
                "company": role[1],
                "title": role[0],
                "start_date": "",
                "end_date": "",
                "description": "",
            }
            continue
        if current and date_match:
            current["start_date"] = _to_iso_date(date_match.group(1))
            current["end_date"] = _to_iso_date(date_match.group(2))
            leftover = DATE_RANGE_RE.sub("", line).strip(" -,|")
            if leftover:
                desc_parts.append(leftover)
            continue
        if current:
            desc_parts.append(line)

    flush()
    return entries[:8]


def _extract_education(text: str) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    for block in re.split(r"\n{2,}", text):
        if not DEGREE_RE.search(block) and "university" not in block.lower() and "college" not in block.lower():
            continue
        line = _clean_block(block)
        degree_match = DEGREE_RE.search(line)
        year_match = re.search(r"(20\d{2}|19\d{2})", line)
        institution = ""
        inst_match = re.search(
            r"([A-Z][A-Za-z&.\s]+(?:University|College|Institute|School))",
            line,
        )
        if inst_match:
            institution = inst_match.group(1).strip()
        field = ""
        field_match = re.search(r"in ([A-Z][A-Za-z&\s]{3,40})", line)
        if field_match:
            field = field_match.group(1).strip()
        entries.append(
            {
                "institution": institution,
                "degree": degree_match.group(0) if degree_match else "",
                "field": field,
                "year": year_match.group(1) if year_match else "",
            }
        )
    return entries[:5]
